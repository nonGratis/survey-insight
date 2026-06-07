"""Export all project code (respecting .gitignore) into a single .docx file.

Formatting follows ДСТУ 3008:2015 (margins 30/10/20/20 mm, Times New Roman 14pt
for body, Courier New 11pt single-spaced for code listings, page numbers
top-right starting from page 2).

Usage:
    pip install python-docx
    python export_code_to_docx.py [-o output.docx] [--strip]

    --strip  remove docstrings and `#` comments from .py files
"""

from __future__ import annotations

import argparse
import ast
import io
import re
import subprocess
import sys
import tokenize
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

EXCLUDE_DIRS = {"research"}
EXCLUDE_EXTS = {".md"}

TEXT_EXTS = {
    ".py",
    ".toml",
    ".txt",
    ".yaml",
    ".yml",
    ".json",
    ".cfg",
    ".ini",
    ".sh",
    ".ps1",
    ".html",
    ".css",
    ".js",
    ".sql",
}
TEXT_NAMES = {
    "Dockerfile",
    "LICENSE",
    ".gitignore",
    ".gitattributes",
    ".dockerignore",
    "requirements.txt",
    "Makefile",
}
ROOT_PRIORITY = ("pyproject.toml", "requirements.txt", "app.py", "Dockerfile")


def is_excluded(rel_path: str) -> bool:
    parts = rel_path.split("/")
    if any(part in EXCLUDE_DIRS for part in parts[:-1]):
        return True
    return Path(rel_path).suffix.lower() in EXCLUDE_EXTS


def list_files(repo_root: Path) -> list[Path]:
    try:
        out = subprocess.check_output(
            ["git", "ls-files", "--cached", "--others", "--exclude-standard"],
            cwd=repo_root,
            text=True,
            encoding="utf-8",
        )
    except (subprocess.CalledProcessError, FileNotFoundError) as e:
        sys.exit(f"ERROR: requires `git` in PATH and a git repo. ({e})")
    return [repo_root / line for line in out.splitlines() if line.strip()]


def is_text(p: Path) -> bool:
    if p.name in TEXT_NAMES:
        return True
    return p.suffix.lower() in TEXT_EXTS


def sort_key(p: Path, repo_root: Path) -> tuple:
    rel = p.relative_to(repo_root).as_posix()
    if rel in ROOT_PRIORITY:
        return (0, ROOT_PRIORITY.index(rel), rel)
    if "/" not in rel:
        return (1, 0, rel)
    return (2, 0, rel)


def read_text(p: Path) -> str | None:
    for enc in ("utf-8", "cp1251"):
        try:
            return p.read_text(encoding=enc)
        except (UnicodeDecodeError, OSError):
            continue
    return None


def strip_py(source: str) -> str:
    """Remove `#` comments and docstrings from a Python source string."""
    try:
        tree = ast.parse(source)
    except SyntaxError:
        return source

    doc_positions: set[tuple[int, int]] = set()
    scoped = (ast.Module, ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)
    for node in ast.walk(tree):
        if isinstance(node, scoped) and node.body and isinstance(node.body[0], ast.Expr):
            v = node.body[0].value
            if isinstance(v, ast.Constant) and isinstance(v.value, str):
                doc_positions.add((v.lineno, v.col_offset))

    try:
        tokens = list(tokenize.generate_tokens(io.StringIO(source).readline))
    except (tokenize.TokenizeError, IndentationError):
        return source

    kept = []
    for tok in tokens:
        if tok.type == tokenize.COMMENT:
            continue
        if tok.type == tokenize.STRING and tok.start in doc_positions:
            continue
        kept.append(tok)

    try:
        result = tokenize.untokenize(kept)
    except (ValueError, IndentationError):
        return source

    result = re.sub(r"\n[ \t]*\n[ \t]*\n+", "\n\n", result)
    return result.lstrip("\n")


def add_page_numbers(section) -> None:
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def setup_dstu(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.different_first_page_header_footer = True
    add_page_numbers(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for lvl in (1, 2):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(14)
        h.font.bold = True
        h.font.color.rgb = None
        hp = h.paragraph_format
        hp.first_line_indent = Cm(0)
        hp.space_before = Pt(12)
        hp.space_after = Pt(6)
        hp.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True


def add_code_block(doc: Document, code: str) -> None:
    for line in code.splitlines() or [""]:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(11)


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("-o", "--output", default="code_export.docx")
    ap.add_argument(
        "--strip",
        action="store_true",
        help="remove docstrings and # comments from .py files",
    )
    args = ap.parse_args()

    repo_root = Path(__file__).resolve().parent
    files = [
        f
        for f in list_files(repo_root)
        if f.is_file() and is_text(f) and not is_excluded(f.relative_to(repo_root).as_posix())
    ]
    files.sort(key=lambda p: sort_key(p, repo_root))

    doc = Document()
    setup_dstu(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    tr = title.add_run("ДОДАТОК А\nЛІСТИНГ ПРОГРАМНОГО КОДУ")
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(14)
    tr.font.bold = True

    info = doc.add_paragraph(
        f"Усього файлів: {len(files)}"
        + ("  (докстрінги та коментарі видалено)" if args.strip else "")
    )
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.first_line_indent = Cm(0)

    doc.add_page_break()

    skipped: list[str] = []
    listed = 0
    for f in files:
        rel = f.relative_to(repo_root).as_posix()
        text = read_text(f)
        if text is None:
            skipped.append(rel)
            continue

        if args.strip and f.suffix.lower() == ".py":
            text = strip_py(text)

        listed += 1
        add_caption(doc, f"Лістинг А.{listed} — {rel}")
        add_code_block(doc, text)

    out_path = repo_root / args.output
    doc.save(out_path)
    print(f"Wrote {out_path} ({listed} files)")
    if skipped:
        print(f"Skipped (unreadable): {len(skipped)}")
        for s in skipped:
            print(f"  - {s}")


if __name__ == "__main__":
    main()
